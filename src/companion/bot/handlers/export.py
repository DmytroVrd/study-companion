import asyncio
import os
import re
import tempfile
from datetime import datetime
from enum import StrEnum
from pathlib import Path
from uuid import uuid4

from aiogram import F, Router
from aiogram.filters import Command
from aiogram.types import (
    CallbackQuery,
    FSInputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    Message,
)
from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession

from companion.db.crud import get_or_create_active_context, get_user_documents
from companion.llm.providers import LLMProviderError, generate_text
from companion.rag.vectorstore import query

router = Router()


class ExportKind(StrEnum):
    SUMMARY = "summary"
    DOCX = "docx"
    ANKI = "anki"
    VOCABULARY = "vocabulary"
    QUIZ = "quiz"


def _detect_export_kind(text: str) -> ExportKind | None:
    normalized = text.lower()
    export_markers = (
        "export",
        "make",
        "create",
        "anki",
        "quiz",
        "test",
        "summary",
        "vocab",
        "flashcard",
        "pdf",
        "docx",
        "word",
        "document",
        "експорт",
        "зроби",
        "створи",
        "тест",
        "конспект",
        "слов",
        "карт",
        "термін",
        "документ",
        "пдф",
        "ворд",
    )
    if not any(marker in normalized for marker in export_markers):
        return None

    document_markers = ("pdf", "docx", "word", "document", "документ", "пдф", "ворд")
    if any(marker in normalized for marker in document_markers):
        return ExportKind.DOCX
    if any(marker in normalized for marker in ("anki", "карт", "flashcard")):
        return ExportKind.ANKI
    if any(marker in normalized for marker in ("слов", "vocab", "термін")):
        return ExportKind.VOCABULARY
    if any(marker in normalized for marker in ("quiz", "test", "тест")):
        return ExportKind.QUIZ
    if any(marker in normalized for marker in ("summary", "конспект", "markdown")):
        return ExportKind.SUMMARY
    return None


def _export_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Markdown summary", callback_data="export:summary"),
                InlineKeyboardButton(text="Word document", callback_data="export:docx"),
            ],
            [
                InlineKeyboardButton(text="Anki CSV", callback_data="export:anki"),
                InlineKeyboardButton(text="Vocabulary", callback_data="export:vocabulary"),
            ],
            [
                InlineKeyboardButton(text="Quiz", callback_data="export:quiz"),
            ],
        ]
    )


def _safe_project_name(title: str) -> str:
    cleaned = re.sub(r"[^\w-]+", "_", title, flags=re.UNICODE).strip("_")
    return cleaned[:48] or "project"


def _new_export_path(extension: str, prefix: str, kind: ExportKind) -> Path:
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    export_id = uuid4().hex[:8]
    filename = f"{prefix}_{kind.value}_{timestamp}_{export_id}{extension}"
    return Path(tempfile.gettempdir()) / filename


def _collect_context(docs: list, request: str) -> str:
    chunks = []
    search_text = (
        f"{request}\nsummary important concepts key terms vocabulary quiz definitions examples"
    )
    for document in docs:
        chunks.extend(query(document.chroma_collection, search_text, k=8))
    return "\n\n".join(chunk.page_content for chunk in chunks)[:12000]


def _prompt_for(kind: ExportKind, context: str) -> tuple[str, str, str]:
    system = (
        "You create study exports from uploaded material. "
        "Use only the provided context. Do not invent facts. "
        "If the material is for learning a language, preserve target-language words "
        "and explain them clearly."
    )
    if kind == ExportKind.ANKI:
        return (
            system,
            "Create 20 Anki-style flashcards from the material. "
            "Return only CSV rows with two columns: Front, Back. "
            "No markdown, no numbering.\n\n"
            f"Material:\n{context}",
            ".csv",
        )
    if kind == ExportKind.VOCABULARY:
        return (
            system,
            "Extract the most useful vocabulary or key terms from the material. "
            "Return markdown with columns: term, meaning, example/context.\n\n"
            f"Material:\n{context}",
            ".md",
        )
    if kind == ExportKind.QUIZ:
        return (
            system,
            "Create a practical quiz from the material: 10 questions, answer key, "
            "and short explanations. Return markdown.\n\n"
            f"Material:\n{context}",
            ".md",
        )
    if kind == ExportKind.DOCX:
        return (
            system,
            "Create a polished study document from the material. Return markdown only. "
            "Include: title, short overview, key ideas, important terms, examples if present, "
            "and 5 review questions. If the user asked for PDF, create content suitable for "
            "a Word document because PDF export is not available yet.\n\n"
            f"Material:\n{context}",
            ".docx",
        )
    return (
        system,
        "Create a clear markdown summary from the material. Include: overview, "
        "key ideas, important terms, and what to review next.\n\n"
        f"Material:\n{context}",
        ".md",
    )


def _fallback_export(kind: ExportKind, context: str) -> str:
    excerpt = context[:2500].strip()
    if kind == ExportKind.ANKI:
        return f"Front,Back\nWhat is this material about?,{excerpt[:400]}"
    return (
        "# Export\n\n"
        "AI providers are unavailable. Here is the most relevant extracted text:\n\n"
        f"{excerpt}"
    )


def _add_markdown_to_docx(document: DocxDocument, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue
        document.add_paragraph(line)


def _write_export_file(content: str, extension: str, prefix: str, kind: ExportKind) -> str:
    export_path = _new_export_path(extension, prefix, kind)
    if extension == ".docx":
        document = DocxDocument()
        document.add_heading(prefix.replace("_", " "), level=0)
        _add_markdown_to_docx(document, content)
        document.save(export_path)
        return str(export_path)

    encoding = "utf-8-sig" if extension == ".csv" else "utf-8"
    with export_path.open("x", encoding=encoding) as tmp:
        tmp.write(content)
        return str(export_path)


async def _send_export(
    message: Message,
    session: AsyncSession,
    kind: ExportKind,
    telegram_id: int,
) -> None:
    docs = await get_user_documents(session, telegram_id)
    if not docs:
        await message.answer("Upload a file first, then ask me to export something.")
        return

    project = await get_or_create_active_context(session, telegram_id)
    if kind == ExportKind.DOCX:
        await message.answer("Preparing Word document...")
    else:
        await message.answer("Preparing export...")

    context = _collect_context(docs, kind.value)
    if len(context.strip()) < 120:
        await message.answer("I could not find enough text in this project to export.")
        return

    system, user, extension = _prompt_for(kind, context)
    try:
        content = await asyncio.to_thread(
            generate_text,
            system=system,
            user=user,
            max_tokens=1800,
            temperature=0.2,
        )
    except LLMProviderError:
        content = _fallback_export(kind, context)

    prefix = _safe_project_name(project.title)
    tmp_name = await asyncio.to_thread(_write_export_file, content, extension, prefix, kind)

    try:
        await message.answer_document(FSInputFile(tmp_name, filename=Path(tmp_name).name))
    finally:
        if os.path.exists(tmp_name):
            os.unlink(tmp_name)


@router.message(Command("export"))
async def cmd_export(message: Message) -> None:
    await message.answer(
        "What do you want to export from the current project?",
        reply_markup=_export_keyboard(),
    )


@router.callback_query(F.data.startswith("export:"))
async def export_callback(callback: CallbackQuery, session: AsyncSession) -> None:
    kind = ExportKind((callback.data or "").split(":", maxsplit=1)[1])
    await _send_export(callback.message, session, kind, callback.from_user.id)
    await callback.answer()


@router.message(F.text.func(lambda text: _detect_export_kind(text or "") is not None))
async def handle_export_intent(message: Message, session: AsyncSession) -> None:
    if (message.text or "").strip().startswith("/"):
        return

    kind = _detect_export_kind(message.text or "")
    if kind is None:
        return
    await _send_export(message, session, kind, message.from_user.id)
