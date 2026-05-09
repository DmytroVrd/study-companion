from pathlib import Path

from docx import Document as DocxDocument

from companion.bot.handlers.export import (
    ExportKind,
    _detect_export_kind,
    _write_export_file,
)


def test_detects_ukrainian_pdf_request_as_docx_export() -> None:
    assert _detect_export_kind("зроби мені з цього щось наприклад pdf") == ExportKind.DOCX


def test_detects_word_document_request_as_docx_export() -> None:
    assert _detect_export_kind("create a word document from this file") == ExportKind.DOCX


def test_writes_docx_export_file() -> None:
    path = _write_export_file(
        "# Title\n\n## Key ideas\n\n- One\n- Two",
        ".docx",
        "test_project",
        ExportKind.DOCX,
    )
    try:
        assert Path(path).exists()
        assert Path(path).suffix == ".docx"
    finally:
        Path(path).unlink(missing_ok=True)


def test_docx_exports_are_new_files_without_appended_content() -> None:
    first_path = Path(
        _write_export_file(
            "# First Export\n\nOnly first content",
            ".docx",
            "test_project",
            ExportKind.DOCX,
        )
    )
    second_path = Path(
        _write_export_file(
            "# Second Export\n\nOnly second content",
            ".docx",
            "test_project",
            ExportKind.DOCX,
        )
    )

    try:
        assert first_path != second_path
        assert first_path.exists()
        assert second_path.exists()

        first_text = "\n".join(p.text for p in DocxDocument(first_path).paragraphs)
        second_text = "\n".join(p.text for p in DocxDocument(second_path).paragraphs)

        assert "Only first content" in first_text
        assert "Only second content" not in first_text
        assert "Only second content" in second_text
        assert "Only first content" not in second_text
    finally:
        first_path.unlink(missing_ok=True)
        second_path.unlink(missing_ok=True)
